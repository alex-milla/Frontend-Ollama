"""
file_exporter.py — Generación de archivos de salida con plantillas.
Formatos: TXT, Markdown, DOCX, PDF, CSV.
"""
import io
import csv
import uuid
import logging
from datetime import date
from pathlib import Path

log = logging.getLogger(__name__)

TEMPLATES = ("informe", "carta", "legal", "email", "resumen", "libre")
FORMATS   = ("txt", "md", "docx", "pdf", "csv")


# ── Dispatch ──────────────────────────────────────────────────────────────────

def generate(text: str, fmt: str, template: str) -> tuple[bytes, str]:
    """
    Devuelve (bytes_del_archivo, mimetype).
    """
    fmt      = fmt.lower()
    template = template.lower()
    if fmt in ("txt", "md"):
        return _gen_txt(text, template, fmt), _mime(fmt)
    elif fmt == "docx":
        return _gen_docx(text, template), _mime("docx")
    elif fmt == "pdf":
        return _gen_pdf(text, template), _mime("pdf")
    elif fmt == "csv":
        return _gen_csv(text), _mime("csv")
    raise ValueError(f"Formato no soportado: {fmt}")


def _mime(fmt: str) -> str:
    return {
        "txt":  "text/plain; charset=utf-8",
        "md":   "text/markdown; charset=utf-8",
        "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "pdf":  "application/pdf",
        "csv":  "text/csv; charset=utf-8",
    }.get(fmt, "application/octet-stream")


def save_output(file_bytes: bytes, fmt: str, display_name: str, project_id: int, output_folder: str) -> tuple[str, str]:
    """
    Guarda en data/outputs/{project_id}/.
    Devuelve (filename_stored, full_path).
    """
    proj_dir = Path(output_folder) / str(project_id)
    proj_dir.mkdir(parents=True, exist_ok=True)
    proj_dir.chmod(0o700)

    base = str(uuid.uuid4())
    filename = base + "." + fmt
    fpath = proj_dir / filename
    fpath.write_bytes(file_bytes)
    fpath.chmod(0o600)
    return filename, str(fpath)


# ── TXT / MD ──────────────────────────────────────────────────────────────────

def _gen_txt(text: str, template: str, fmt: str) -> bytes:
    today = date.today().strftime("%d/%m/%Y")
    header = _txt_header(template, today)
    return (header + text).encode("utf-8")


def _txt_header(template: str, today: str) -> str:
    if template == "informe":
        return f"INFORME\nFecha: {today}\n{'='*60}\n\n"
    elif template == "carta":
        return f"{today}\n\nEstimado/a señor/a:\n\n"
    elif template == "legal":
        return f"DOCUMENTO LEGAL\nFecha: {today}\n\nEXPONE:\n\n"
    elif template == "email":
        return f"Para: \nCC: \nAsunto: \nFecha: {today}\n\n"
    elif template == "resumen":
        return f"RESUMEN EJECUTIVO — {today}\n{'='*60}\n\n"
    return ""  # libre


# ── DOCX ──────────────────────────────────────────────────────────────────────

def _gen_docx(text: str, template: str) -> bytes:
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Márgenes estándar
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3)
        section.right_margin  = Cm(2.5)

    today = date.today().strftime("%d de %B de %Y")

    if template == "informe":
        _docx_informe(doc, text, today)
    elif template == "carta":
        _docx_carta(doc, text, today)
    elif template == "legal":
        _docx_legal(doc, text, today)
    elif template == "email":
        _docx_email(doc, text, today)
    elif template == "resumen":
        _docx_resumen(doc, text, today)
    else:
        _docx_libre(doc, text)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _set_style(para, size_pt=11, bold=False, align=None):
    from docx.shared import Pt
    run = para.runs[0] if para.runs else para.add_run()
    run.font.size = Pt(size_pt)
    run.bold = bold
    if align:
        para.alignment = align


def _docx_informe(doc, text, today):
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    p = doc.add_heading("INFORME", level=1)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Fecha: {today}").alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc.add_paragraph()
    _add_body(doc, text)
    _add_page_numbers(doc)


def _docx_carta(doc, text, today):
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    doc.add_paragraph(today).alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc.add_paragraph()
    doc.add_paragraph("Estimado/a señor/a:")
    doc.add_paragraph()
    _add_body(doc, text)
    doc.add_paragraph()
    doc.add_paragraph("Atentamente,")
    doc.add_paragraph()
    doc.add_paragraph("_________________________")


def _docx_legal(doc, text, today):
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    p = doc.add_heading("DOCUMENTO LEGAL", level=1)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Fecha: {today}").alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc.add_paragraph()
    doc.add_heading("EXPONE:", level=2)
    _add_body(doc, text)
    doc.add_paragraph()
    doc.add_heading("SOLICITA:", level=2)
    doc.add_paragraph("Lo que corresponda conforme a derecho.")
    doc.add_paragraph()
    doc.add_paragraph(f"En _____________, a {today}")


def _docx_email(doc, text, today):
    doc.add_heading("CORREO ELECTRÓNICO", level=1)
    for label in ("Para:", "CC:", "Asunto:"):
        p = doc.add_paragraph()
        p.add_run(label + " ").bold = True
    doc.add_paragraph(f"Fecha: {today}")
    doc.add_paragraph()
    _add_body(doc, text)


def _docx_resumen(doc, text, today):
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    p = doc.add_heading("RESUMEN EJECUTIVO", level=1)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(today).alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc.add_paragraph()
    doc.add_heading("Puntos clave:", level=2)
    for line in text.split("\n"):
        line = line.strip()
        if line:
            doc.add_paragraph(line, style="List Bullet")
    doc.add_paragraph()
    doc.add_heading("Conclusión:", level=2)
    doc.add_paragraph("Ver desarrollo completo arriba.")


def _docx_libre(doc, text):
    _add_body(doc, text)


def _add_body(doc, text):
    for line in text.split("\n"):
        doc.add_paragraph(line)


def _add_page_numbers(doc):
    # Footer con número de página (via XML nativo)
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    for section in doc.sections:
        footer = section.footer
        para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        para.clear()
        run = para.add_run()
        fldChar1 = OxmlElement("w:fldChar"); fldChar1.set(qn("w:fldCharType"), "begin")
        instrText = OxmlElement("w:instrText"); instrText.text = "PAGE"; instrText.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        fldChar2 = OxmlElement("w:fldChar"); fldChar2.set(qn("w:fldCharType"), "end")
        run._r.append(fldChar1); run._r.append(instrText); run._r.append(fldChar2)


# ── PDF ───────────────────────────────────────────────────────────────────────

def _gen_pdf(text: str, template: str) -> bytes:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import cm
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.enums import TA_CENTER, TA_RIGHT, TA_JUSTIFY
    from reportlab.lib import colors
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
    from reportlab.platypus import PageTemplate, Frame
    from reportlab.pdfgen import canvas as rl_canvas

    buf = io.BytesIO()
    today = date.today().strftime("%d/%m/%Y")

    styles = getSampleStyleSheet()
    base_style = ParagraphStyle(
        "base", parent=styles["Normal"],
        fontName="Helvetica", fontSize=10, leading=15,
        leftIndent=0, rightIndent=0, spaceAfter=6,
    )
    title_style = ParagraphStyle(
        "title", parent=base_style,
        fontSize=18, fontName="Helvetica-Bold",
        alignment=TA_CENTER, spaceAfter=12,
    )
    h2_style = ParagraphStyle(
        "h2", parent=base_style,
        fontSize=12, fontName="Helvetica-Bold", spaceAfter=6,
    )
    right_style = ParagraphStyle(
        "right", parent=base_style, alignment=TA_RIGHT,
    )

    story = []

    def add_text(t):
        for line in t.split("\n"):
            story.append(Paragraph(line or "&nbsp;", base_style))

    if template == "informe":
        story.append(Paragraph("INFORME", title_style))
        story.append(Paragraph(f"Fecha: {today}", right_style))
        story.append(HRFlowable(width="100%", thickness=1, color=colors.grey))
        story.append(Spacer(1, 0.4*cm))
        add_text(text)

    elif template == "carta":
        story.append(Paragraph(today, right_style))
        story.append(Spacer(1, 0.4*cm))
        story.append(Paragraph("Estimado/a señor/a:", base_style))
        story.append(Spacer(1, 0.3*cm))
        add_text(text)
        story.append(Spacer(1, 0.5*cm))
        story.append(Paragraph("Atentamente,", base_style))
        story.append(Spacer(1, 1*cm))
        story.append(Paragraph("_________________________", base_style))

    elif template == "legal":
        story.append(Paragraph("DOCUMENTO LEGAL", title_style))
        story.append(Paragraph(f"Fecha: {today}", right_style))
        story.append(HRFlowable(width="100%", thickness=1, color=colors.grey))
        story.append(Spacer(1, 0.4*cm))
        story.append(Paragraph("EXPONE:", h2_style))
        add_text(text)
        story.append(Spacer(1, 0.5*cm))
        story.append(Paragraph("SOLICITA:", h2_style))
        story.append(Paragraph("Lo que corresponda conforme a derecho.", base_style))
        story.append(Spacer(1, 0.5*cm))
        story.append(Paragraph(f"En _____________, a {today}", base_style))

    elif template == "email":
        story.append(Paragraph("CORREO ELECTRÓNICO", title_style))
        for label in ("Para:", "CC:", "Asunto:"):
            story.append(Paragraph(f"<b>{label}</b> ____________________", base_style))
        story.append(Paragraph(f"Fecha: {today}", base_style))
        story.append(HRFlowable(width="100%", thickness=0.5, color=colors.grey))
        story.append(Spacer(1, 0.3*cm))
        add_text(text)

    elif template == "resumen":
        story.append(Paragraph("RESUMEN EJECUTIVO", title_style))
        story.append(Paragraph(today, right_style))
        story.append(HRFlowable(width="100%", thickness=1, color=colors.grey))
        story.append(Spacer(1, 0.4*cm))
        story.append(Paragraph("Puntos clave:", h2_style))
        for line in text.split("\n"):
            line = line.strip()
            if line:
                story.append(Paragraph(f"• {line}", base_style))
        story.append(Spacer(1, 0.3*cm))
        story.append(Paragraph("Conclusión:", h2_style))
        story.append(Paragraph("Ver puntos clave arriba.", base_style))

    else:  # libre
        add_text(text)

    def on_page(c, doc):
        c.saveState()
        c.setFont("Helvetica", 8)
        c.setFillColor(colors.grey)
        c.drawRightString(A4[0] - 2*cm, 1.2*cm, f"Página {doc.page}")
        c.restoreState()

    doc = SimpleDocTemplate(
        buf, pagesize=A4,
        rightMargin=2.5*cm, leftMargin=3*cm,
        topMargin=2.5*cm, bottomMargin=2.5*cm,
    )
    doc.build(story, onFirstPage=on_page, onLaterPages=on_page)
    return buf.getvalue()


# ── CSV ───────────────────────────────────────────────────────────────────────

def _gen_csv(text: str) -> bytes:
    buf = io.StringIO()
    writer = csv.writer(buf)
    for line in text.split("\n"):
        writer.writerow([line])
    return buf.getvalue().encode("utf-8")
